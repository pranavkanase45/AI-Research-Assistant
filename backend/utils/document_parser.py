"""
Unified Document Parser
Supports: PDF, DOCX, TXT, HTML
"""
import os
from typing import Tuple
import pdfplumber
from docx import Document
from bs4 import BeautifulSoup
from utils.logger import parser_logger


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file

    Args:
        file_path: Path to PDF file

    Returns:
        Extracted text content
    """
    parser_logger.info(f"Extracting text from PDF: {file_path}")
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            num_pages = len(pdf.pages)
            parser_logger.debug(f"PDF has {num_pages} pages")
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                text += page_text
                parser_logger.debug(f"Extracted {len(page_text)} chars from page {i}/{num_pages}")
        parser_logger.info(f"PDF extraction complete: {len(text)} total characters")
    except Exception as e:
        parser_logger.error(f"Error extracting PDF: {str(e)}")
        raise ValueError(f"Error extracting PDF: {str(e)}")
    return text


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file

    Args:
        file_path: Path to DOCX file

    Returns:
        Extracted text content
    """
    parser_logger.info(f"Extracting text from DOCX: {file_path}")
    text = ""
    try:
        doc = Document(file_path)

        # Extract paragraphs
        parser_logger.debug(f"Extracting {len(doc.paragraphs)} paragraphs")
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Extract tables
        parser_logger.debug(f"Extracting {len(doc.tables)} tables")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"

        parser_logger.info(f"DOCX extraction complete: {len(text)} total characters")
    except Exception as e:
        parser_logger.error(f"Error extracting DOCX: {str(e)}")
        raise ValueError(f"Error extracting DOCX: {str(e)}")
    return text


def extract_text_from_html(file_path: str) -> str:
    """
    Extract text from HTML file

    Args:
        file_path: Path to HTML file

    Returns:
        Extracted text content
    """
    parser_logger.info(f"Extracting text from HTML: {file_path}")
    text = ""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()

        parser_logger.debug(f"HTML file size: {len(html_content)} characters")

        # Parse HTML
        soup = BeautifulSoup(html_content, 'lxml')

        # Remove script and style elements
        removed_elements = 0
        for script in soup(["script", "style", "meta", "link"]):
            script.decompose()
            removed_elements += 1
        parser_logger.debug(f"Removed {removed_elements} non-text elements")

        # Get text
        text = soup.get_text(separator='\n', strip=True)

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)

        parser_logger.info(f"HTML extraction complete: {len(text)} total characters")
    except Exception as e:
        parser_logger.error(f"Error extracting HTML: {str(e)}")
        raise ValueError(f"Error extracting HTML: {str(e)}")
    return text


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from TXT file

    Args:
        file_path: Path to TXT file

    Returns:
        Extracted text content
    """
    parser_logger.info(f"Extracting text from TXT: {file_path}")
    text = ""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        parser_logger.info(f"TXT extraction complete: {len(text)} total characters")
    except Exception as e:
        parser_logger.error(f"Error extracting TXT: {str(e)}")
        raise ValueError(f"Error extracting TXT: {str(e)}")
    return text


def extract_text_from_file(file_path: str) -> Tuple[str, str]:
    """
    Extract text from any supported file format
    Automatically detects format based on file extension

    Supported formats:
    - PDF (.pdf)
    - DOCX (.docx)
    - HTML (.html, .htm)
    - TXT (.txt)

    Args:
        file_path: Path to file

    Returns:
        Tuple of (extracted_text, file_type)

    Raises:
        ValueError: If file format is not supported or extraction fails
    """
    # Get file extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    parser_logger.info(f"Processing file: {file_path} (extension: {ext})")

    # Map extensions to parsers
    parsers = {
        '.pdf': (extract_text_from_pdf, 'PDF'),
        '.docx': (extract_text_from_docx, 'DOCX'),
        '.html': (extract_text_from_html, 'HTML'),
        '.htm': (extract_text_from_html, 'HTML'),
        '.txt': (extract_text_from_txt, 'TXT'),
    }

    if ext not in parsers:
        supported = ', '.join(parsers.keys())
        parser_logger.warning(f"Unsupported file format: {ext}")
        raise ValueError(f"Unsupported file format: {ext}. Supported formats: {supported}")

    # Extract text using appropriate parser
    parser_func, file_type = parsers[ext]
    parser_logger.debug(f"Using parser for {file_type}")
    text = parser_func(file_path)

    if not text or len(text.strip()) < 10:
        parser_logger.warning(f"No meaningful text extracted from {file_type} file")
        raise ValueError(f"No meaningful text extracted from {file_type} file")

    parser_logger.info(f"File processed successfully: {file_type}, {len(text)} characters")
    return text, file_type


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """
    Split text into overlapping chunks

    Args:
        text: Text to chunk
        chunk_size: Size of each chunk in characters
        overlap: Number of overlapping characters between chunks

    Returns:
        List of text chunks
    """
    if not text:
        parser_logger.warning("Empty text provided for chunking")
        return []

    parser_logger.debug(f"Chunking text: length={len(text)}, chunk_size={chunk_size}, overlap={overlap}")

    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]

        # Only add non-empty chunks
        if chunk.strip():
            chunks.append(chunk)

        start += (chunk_size - overlap)

    parser_logger.info(f"Text chunked into {len(chunks)} chunks")
    return chunks


# Supported file extensions
SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.html', '.htm', '.txt']
SUPPORTED_MIME_TYPES = [
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'text/html',
    'text/plain'
]
